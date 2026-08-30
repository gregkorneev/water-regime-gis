from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


PATH = "output/docx/Тезисы_оценка_влажности_почвы_по_ДЗЗ_и_модели.docx"
TEXT = (
    "Архитектура информационной системы. При реализации информационной системы "
    "для распределения геоинформационных данных использован QGIS. Основу составляют "
    "локальные геопространственные данные. Для выполнения расчётов разработан модуль "
    "на Python, позволяющий загружать контуры полей, временные ряды из внешнего сервиса "
    "и снимки Sentinel-1 и Sentinel-2 за заданный интервал с фильтрацией облачных сцен. "
    "Растровые, векторные и табличные данные хранятся локально и объединяются проектом "
    "QGIS. По спутниковым снимкам рассчитываются индексы. Скрипт сопоставляет временные "
    "ряды внешней модели с динамикой индексов и подготавливает данные для построения "
    "моделей, описывающих их связь. Система автоматизирована: после поступления новых "
    "внешних данных она докачивает спутниковые данные, завершает расчёты и обновляет графики."
)


def set_font(run, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = None
    run.bold = bold


document = Document(PATH)
paragraph = document.paragraphs[9]
paragraph.clear()
heading = paragraph.add_run("Архитектура информационной системы. ")
set_font(heading, bold=True)
body = paragraph.add_run(TEXT.split(". ", 1)[1])
set_font(body)
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
document.save(PATH)
