from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

OUT = Path("output/docx/Тезисы_оценка_влажности_почвы_по_ДЗЗ_и_модели.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2)
sec.right_margin = Cm(2)
sec.header_distance = Cm(1.25)
sec.footer_distance = Cm(1.25)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.font.size = Pt(12)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_before = Pt(0)
pf.space_after = Pt(0)

for name in ("Heading 1", "Heading 2", "Heading 3"):
    st = styles[name]
    st.font.name = "Times New Roman"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    st.font.size = Pt(12)
    st.font.bold = True
    st.font.color.rgb = None
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.space_after = Pt(0)

def set_run(run, *, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic

def p(text="", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, before=0, after=0):
    para = doc.add_paragraph()
    para.alignment = align
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if first_line:
        fmt.first_line_indent = Cm(1.25)
    r = para.add_run(text)
    set_run(r)
    return para

def label_paragraph(label, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.first_line_indent = Cm(1.25)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    r = para.add_run(label)
    set_run(r, bold=True)
    r = para.add_run(text)
    set_run(r)
    return para

# Conference template: UDC, title, author details, abstract and keywords.
p("УДК 004.6:631.67", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
title = p("ОЦЕНКА ВЛАЖНОСТИ ПОЧВЫ НА СЕЛЬСКОХОЗЯЙСТВЕННЫХ ПОЛЯХ ПО СПУТНИКОВЫМ ДАННЫМ И ФИЗИЧЕСКИ ОБОСНОВАННОЙ МОДЕЛИ", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
for run in title.runs:
    run.bold = True

author = p("[Фамилия Имя Отчество автора]", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
for run in author.runs:
    run.bold = True
p("[Организация, город, Россия]", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
p("[e-mail автора]", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

label_paragraph("Аннотация. ", "Предложен подход к оценке пространственно распределённой влажности почвы на сельскохозяйственных полях на основе точечных наземных измерений, продуктов Sentinel-2 и физически обоснованного моделирования. Для вегетационного периода 2026 года рассчитываются индексы NDVI, NDWI и анализируется продукт FCOVER. Спутниковые показатели сопоставляются с модельными характеристиками развития растительного покрова и запасов влаги. Подход обеспечивает переход от локальных показаний датчиков к оценке состояния всего поля для управления поливами.")
label_paragraph("Ключевые слова: ", "влажность почвы, дистанционное зондирование Земли, Sentinel-2, NDVI, NDWI, FCOVER, орошение, математическое моделирование.")

label_paragraph("Введение. ", "Эффективное управление поливами требует своевременной оценки влагообеспеченности посевов. Датчики влажности характеризуют лишь отдельные точки и не отражают неоднородность поля, обусловленную свойствами почв, микрорельефом и состоянием растительного покрова. Поэтому актуальна задача восстановления пространственной картины влажности по данным для всей площади поля.")
p("Спутниковые данные Sentinel-2 с разрешением до 10 м позволяют наблюдать развитие посевов. Индексы рассчитываются по отражательной способности поверхности: NDVI = (ρNIR − ρRED) / (ρNIR + ρRED), где для Sentinel-2 используются канал B8 ближнего инфракрасного диапазона (NIR, 10 м) и канал B4 красного диапазона (RED, 10 м). Рост NDVI обычно соответствует увеличению доли фотосинтезирующей биомассы; низкие значения характерны для открытой почвы, а отрицательные — для воды, облаков или теней.")
p("Для оценки водного состояния растительного покрова рассчитывается NDWI = (ρNIR − ρSWIR) / (ρNIR + ρSWIR), где ρSWIR — отражательная способность канала B11 коротковолнового инфракрасного диапазона; перед расчётом B11 приводится к разрешению 10 м. Более высокие значения NDWI указывают на большее содержание влаги в растительности, однако индекс следует интерпретировать совместно с NDVI и метеоданными. FCOVER — безразмерный спутниковый продукт от 0 до 1, характеризующий долю площади пикселя, занятую зелёной растительностью. Его сопоставление с модельной долей покрытия поля позволяет независимо проверить динамику развития посевов [1, 2].")

label_paragraph("Материалы и методы. ", "Исследование выполняется для рассматриваемых полей с апреля по август 2026 года. Используются измерения датчиков, сцены Sentinel-2, продукт поверхностной влажности по Sentinel-1, границы полей и метеорологические данные. Обработка включает отбор безоблачных сцен, маскирование облаков и расчёт NDVI, NDWI и FCOVER.")
p("Модель водного баланса рассчитывает развитие растений и влажность почвы с учётом осадков, поливов, испарения, транспирации и запасов влаги. Значения поверхностной влажности Sentinel-1 для слоя 5-10 см сопоставляются с расчётной влажностью той же глубины; учитывается влияние растительности и шероховатости на радиолокационный сигнал [3]. Модель проверяется по наземным измерениям с использованием MAE, RMSE и R².")

label_paragraph("Результаты и обсуждение. ", "Получены временные ряды NDVI, NDWI, FCOVER, поверхностной влажности Sentinel-1 и модельных характеристик. FCOVER сопоставляется с модельной долей покрытия поля, а Sentinel-1 — с модельной влажностью верхнего слоя, что дополняет проверку результатов между точками датчиков. В период активной вегетации рост FCOVER и NDVI сопровождается изменением модельной транспирации и влагозапаса.")
p("Совмещение спутниковых и модельных показателей позволяет формировать карты влагообеспеченности поля; в окончательную версию включаются MAE, RMSE, R² и показатели связи индексов с модельной динамикой растительности.")

label_paragraph("Заключение. ", "Для оценки влажности глубже верхнего слоя почвы необходима физически обоснованная модель, входами которой являются погодные данные и свойства почвы. Параметры развития растительности задаются по FCOVER с точностью, оцениваемой сопоставлением с модельной долей покрытия. Совмещение Sentinel-1, Sentinel-2, наземных измерений и модели обеспечивает пространственную оценку состояния поля для управления поливами.")

p("Список источников", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, before=6)
doc.paragraphs[-1].runs[0].bold = True

references = [
    "1. Drusch M., Del Bello U., Carlier S. et al. Sentinel-2: ESA’s Optical High-Resolution Mission for GMES Operational Services. Remote Sensing of Environment. 2012. Vol. 120. P. 25–36.",
    "2. Copernicus Global Land Service. FCOVER: Fraction of Green Vegetation Cover. Product User Manual. European Union, 2024.",
    "3. Bauer-Marschallinger B. et al. Toward Global Soil Moisture Monitoring with Sentinel-1. IEEE Transactions on Geoscience and Remote Sensing. 2019. Vol. 57, no. 1. P. 520-539.",
]
for item in references:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.left_indent = Cm(0.75)
    fmt.first_line_indent = Cm(-0.75)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    set_run(para.add_run(item))

doc.core_properties.title = "Оценка влажности почвы на сельскохозяйственных полях"
doc.core_properties.author = "[Фамилия Имя Отчество автора]"
doc.save(OUT)
print(OUT)
